from PyQt6.QtCore import QThread, pyqtSignal
import os
from src.core.ocr_engine import OCREngine
from docx import Document 

class ExportWorker(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    
    def __init__(self, pdf_handler, page_indices, output_dir, fmt, dpi):
        super().__init__()
        self.pdf_handler = pdf_handler
        self.page_indices = page_indices
        self.output_dir = output_dir
        self.fmt = fmt.lower()
        self.dpi = int(dpi)

    def run(self):
        total = len(self.page_indices)
        for i, page_index in enumerate(self.page_indices):
            filename = f"page_{page_index + 1}.{self.fmt}"
            full_path = os.path.join(self.output_dir, filename)
            self.pdf_handler.save_page(page_index, full_path, self.dpi, self.fmt)
            percent = int((i + 1) / total * 100)
            self.progress.emit(percent)
        self.finished.emit("Export Complete!")

class OCRWorker(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)

    # Added 'lang' parameter here
    def __init__(self, pdf_handler, page_indices, output_path, is_merge, file_type, lang="eng"):
        super().__init__()
        self.pdf_handler = pdf_handler
        self.page_indices = page_indices
        self.output_path = output_path
        self.is_merge = is_merge
        self.file_type = file_type
        self.lang = lang # Store the language
        self.ocr_engine = OCREngine()

    def run(self):
        total = len(self.page_indices)
        merged_text = ""          
        doc = Document()          
        
        temp_image_path = "temp_ocr_processing.png"
        
        for i, page_index in enumerate(self.page_indices):
            self.pdf_handler.save_page(page_index, temp_image_path, dpi=300, fmt="png")
            
            # Pass the specific language to the engine
            text = self.ocr_engine.extract_text(temp_image_path, lang=self.lang)
            
            if self.is_merge:
                if self.file_type == "txt":
                    merged_text += f"--- Page {page_index + 1} ---\n{text}\n\n"
                else: 
                    doc.add_heading(f'Page {page_index + 1}', level=2)
                    doc.add_paragraph(text)
                    doc.add_page_break()
            else:
                if self.file_type == "txt":
                    fname = os.path.join(self.output_path, f"page_{page_index + 1}.txt")
                    with open(fname, "w", encoding="utf-8") as f:
                        f.write(text)
                else: 
                    fname = os.path.join(self.output_path, f"page_{page_index + 1}.docx")
                    single_doc = Document()
                    single_doc.add_paragraph(text)
                    single_doc.save(fname)

            percent = int((i + 1) / total * 100)
            self.progress.emit(percent)

        if self.is_merge:
            if self.file_type == "txt":
                with open(self.output_path, "w", encoding="utf-8") as f:
                    f.write(merged_text)
            else: 
                doc.save(self.output_path)

        if os.path.exists(temp_image_path):
            os.remove(temp_image_path)
            
        self.finished.emit("OCR Processing Complete!")